import os
import fitz  # PyMuPDF
from docx import Document

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text content from a PDF file using PyMuPDF."""
    try:
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            page_text = page.get_text()
            if page_text:
                text += page_text + "\n"
        doc.close()
        return text.strip()
    except Exception as e:
        raise ValueError(f"Failed to parse PDF file using PyMuPDF: {str(e)}")

def extract_text_from_docx(file_path: str) -> str:
    """Extract text content from a DOCX file, including paragraphs and tables."""
    try:
        doc = Document(file_path)
        text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            text.append(paragraph.text.strip())
        return "\n".join(text).strip()
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX file: {str(e)}")

def extract_text(file_path: str) -> str:
    """Detect file extension and extract text accordingly."""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found at path: {file_path}")
    
    _, ext = os.path.splitext(file_path.lower())
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        return extract_text_from_docx(file_path)
    elif ext == ".txt":
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read().strip()
    else:
        raise ValueError(f"Unsupported file format: {ext}. Only PDF, DOCX, and TXT are supported.")
